import io

from docx import Document
from openpyxl import Workbook

from app.services.documents.parsers.csv_parser import parse_csv
from app.services.documents.parsers.docx_parser import parse_docx
from app.services.documents.parsers.pdf_parser import parse_pdf
from app.services.documents.parsers.txt_parser import parse_txt
from app.services.documents.parsers.xlsx_parser import parse_xlsx


def _sample_pdf_bytes() -> bytes:
    from pathlib import Path

    path = Path(__file__).resolve().parents[3] / "sample_data" / "sample_contract.pdf"
    return path.read_bytes()


def test_pdf_parser_preserves_page_numbers():
    segments, page_count = parse_pdf(_sample_pdf_bytes())
    assert page_count == 2
    assert len(segments) == 2
    assert segments[0].page_number == 1
    assert segments[1].page_number == 2
    assert "APPROVED CONTRACT VALUE" in segments[1].text


def test_docx_parser_preserves_headings():
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the intro paragraph.")
    doc.add_heading("Terms", level=1)
    doc.add_paragraph("This is the terms paragraph.")
    buf = io.BytesIO()
    doc.save(buf)

    segments = parse_docx(buf.getvalue())
    assert len(segments) == 2
    assert segments[0].section_title == "Introduction"
    assert "intro paragraph" in segments[0].text
    assert segments[1].section_title == "Terms"
    assert "terms paragraph" in segments[1].text


def test_docx_parser_handles_no_headings():
    doc = Document()
    doc.add_paragraph("Just a plain paragraph with no heading.")
    buf = io.BytesIO()
    doc.save(buf)

    segments = parse_docx(buf.getvalue())
    assert len(segments) == 1
    assert segments[0].section_title is None


def test_xlsx_parser_preserves_sheet_name_and_row_ranges():
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoices"
    ws.append(["id", "value"])
    for i in range(5):
        ws.append([i, i * 10])
    buf = io.BytesIO()
    wb.save(buf)

    segments = parse_xlsx(buf.getvalue())
    assert len(segments) == 1
    assert "Invoices" in segments[0].section_title
    assert "id=0" in segments[0].text


def test_xlsx_parser_never_produces_one_mega_chunk_for_large_sheet():
    wb = Workbook()
    ws = wb.active
    ws.append(["id"])
    for i in range(150):
        ws.append([i])
    buf = io.BytesIO()
    wb.save(buf)

    segments = parse_xlsx(buf.getvalue())
    assert len(segments) > 1  # 150 rows at 50/chunk -> 3 segments


def test_csv_parser_preserves_row_ranges():
    csv_bytes = b"id,value\n1,10\n2,20\n3,30\n"
    segments = parse_csv(csv_bytes)
    assert len(segments) == 1
    assert "rows 1-3" in segments[0].section_title
    assert "id=1" in segments[0].text


def test_txt_parser_returns_single_segment():
    segments = parse_txt(b"just some plain text content")
    assert len(segments) == 1
    assert segments[0].page_number is None
    assert segments[0].section_title is None


def test_txt_parser_empty_file_returns_no_segments():
    assert parse_txt(b"") == []
    assert parse_txt(b"   ") == []
